# IMPORTANT: 1. Change the VPN channel to a fixed one instead of auto one such as 全球智能线路.
# 2. Close the antivirus program if possible.
# 3. Try changing the .Request browser if no other errors found.
# 4. ValueError -> !the error caused by directly extracting some non-English chars to Word, can extract to .txt first
# 5. Extremely recommend using .get_text() to extract human readable contents

import urllib.request, urllib.parse, urllib.error
from bs4 import BeautifulSoup
import ssl

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from wxc_url import indexDict

import datetime

print(datetime.datetime.now())  # Start run-time

# to avoid no cert error
ctx = ssl.create_default_context()
ctx.check_hostname = False
ctx.verify_mode = ssl.CERT_NONE

# indexDict = indexDict()               # Uncommented when testing
# for url in indexDict.values():        # Uncommented when testing

for url in ["https://blog.wenxuecity.com/myblog/1666/201411/16002.html"]:         # for testing only

    browser = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.77 Safari/537.36'
    page = urllib.request.Request(url, headers={'User-Agent': browser})
    fhand = urllib.request.urlopen(page, context=ctx).read().decode('UTF-8')
    # print(fhand)      # for testing only

    doc = Document()  # Open a NEW word document
    # Set normal styles for the whole document**
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  # Independently change the Chinese font name!!*
    font.size = Pt(12)

    soup = BeautifulSoup(fhand, 'html.parser')  # soup is a class object

    title = soup.find('h2').get_text()  # soup.find('x') will only find the first 'x'(use .find_all('x') if all)
    doc.add_heading(title, 0)

    tags = soup('span')  # convert soup to a .ResultSet of lines with tag 'span' (similar with a list, can use .len() method)
    date = ''
    for tag in tags:  # read line by line - tag is .tag type
        if 'class' in tag.attrs and tag['class'][0] == 'time':  # ! <span class="time BLK_txtc"> has ['time', 'BLK_txtc'] as class value!
            date = tag.get_text()[1:11]
            doc.add_paragraph(date)
            break

    tags = soup.find('div', class_="articalContent")  # convert soup to a .ResultSet of lines with tag 'p' (similar with a list, can use .len() method)
    # The following is a Great way of find human-readable texts and break same with original look
    text = tags.get_text('|', strip=True)       # only human-readable texts inside a tag with join mark & strip each
    lines = text.split('|')
    for line in lines:
        try:        # Avoid ValueError: All strings must be XML compatible: Unicode or ASCII, no NULL bytes
            # or control characters; !the error caused due to directly extracting to Word, can extract to .txt first
            doc.add_paragraph(line)
        except:
            continue    # skip the line

    doc.add_paragraph("======================================================")

    tags = soup('div', attrs={'class': 'comment-r'})  # convert soup to a .ResultSet of lines with tag 'div' and class = 'comment-r'

    for tag in tags:
        if tag.contents[1].contents[1].contents == ['润涛阎']:
            doc.add_paragraph('润涛阎 comment:')
            for line in tag.contents[9].contents:
                try:
                    doc.add_paragraph(line.strip())
                except:
                    continue
            doc.add_paragraph('---------------------------------------------------------------------------------')

    # Avoid illegal file name error
    title = title.replace('\\', '').replace('/', '').replace(':', '').replace('*', '').replace('?', '')\
        .replace('"', '').replace('<', '').replace('>', '').replace('|', '')

    path = f"C:\\Users\\MINRUI\\Desktop\\reWXC\\{title}({date}).docx"
    doc.save(path)

print(datetime.datetime.now())  # End run-time

# Old codes with logic errors: find human-readable contents (lines 54-58)
# tags = soup('p')  # convert soup to a .ResultSet of lines with tag 'p' (similar with a list, can use .len() method)
# for tag in tags:  # read line by line - *tag is .tag type
#     if 'class' not in tag.attrs:  # .tag is a special main object in BeautifulSoup
#         doc.add_paragraph(tag.get_text())