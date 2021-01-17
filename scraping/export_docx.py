import urllib.request,urllib.parse,urllib.error
from bs4 import BeautifulSoup
import ssl

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

# to avoid no cert error
ctx = ssl.create_default_context()
ctx.check_hostname=False
ctx.verify_mode=ssl.CERT_NONE

url = "https://blog.wenxuecity.com/myblog/1666/202010/25516.html"
browser = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.77 Safari/537.36'
page = urllib.request.Request(url, headers={'User-Agent': browser})
fhand = urllib.request.urlopen(page, context=ctx).read().decode('UTF-8')
# print(fhand)

doc = Document()
# Set normal styles for the whole document**
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')   # Independently change the Chinese font name!!*
font.size = Pt(12)

soup = BeautifulSoup(fhand, 'html.parser')  # soup is a class object

title = soup.find('h2').get_text()   # soup.find('x') will only find the first 'x'(use .find_all('x') if all)
doc.add_heading(title, 0)

tags = soup('span')    # convert soup to a .ResultSet of lines with tag 'span' (similar with a list, can use .len() method)
date = ''
for tag in tags:    # read line by line - tag is .tag type
    if 'class' in tag.attrs and tag['class'][0] == 'time':    #! <span class="time BLK_txtc"> has ['time', 'BLK_txtc'] as class value!
        date = tag.get_text()[1:11]
        doc.add_paragraph(date)
        break

tags = soup('p')    # convert soup to a .ResultSet of lines with tag 'p' (similar with a list, can use .len() method)
for tag in tags:    # read line by line - *tag is .tag type
    if 'class' not in tag.attrs:    # .tag is a special main object in BeautifulSoup
        doc.add_paragraph(tag.get_text())

doc.add_paragraph("======================================================")


tags = soup('div', attrs={'class': 'comment-r'})    # convert soup to a .ResultSet of lines with tag 'div' and class = 'comment-r'

for tag in tags:
    if tag.contents[1].contents[1].contents == ['润涛阎']:
        doc.add_paragraph('润涛阎 comment:')
        for line in tag.contents[9].contents:
            try:
                doc.add_paragraph(line.strip())
            except:
                continue
        doc.add_paragraph('------------------------------------------------------------------')


path = f"C:\\Users\\MINRUI\\Desktop\\WXC\\{title}({date}).docx"
doc.save(path)